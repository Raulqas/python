import tkinter as tk
from tkinter import messagebox
from tkinter.filedialog import asksaveasfilename
from abc import ABC, abstractmethod
from docx import Document
from openpyxl import Workbook


class Shape(ABC):
    @abstractmethod
    def calculate_volume(self):
        pass

    @abstractmethod
    def calculate_surface_area(self):
        pass

    @abstractmethod
    def calculate_mass(self, volume, density):
        pass


class Parallelepiped(Shape):
    def __init__(self, length, width, height):
        self.length = length
        self.width = width
        self.height = height

    def calculate_volume(self):
        return self.length * self.width * self.height

    def calculate_surface_area(self):
        return 2 * (self.length * self.width + self.width * self.height + self.height * self.length)

    def calculate_mass(self, volume, density):
        return volume * density


class Tetrahedron(Shape):
    def __init__(self, length, height):
        self.length = length
        self.height = height

    def calculate_volume(self):
        return (self.length ** 2) * self.height / (6 * (2 ** 0.5))

    def calculate_surface_area(self):
        return self.length ** 2 * (3 ** 0.5)

    def calculate_mass(self, volume, density):
        return volume * density


class Sphere(Shape):
    def __init__(self, radius):
        self.radius = radius

    def calculate_volume(self):
        return (4 / 3) * 3.14 * (self.radius ** 3)

    def calculate_surface_area(self):
        return 4 * 3.14 * (self.radius ** 2)

    def calculate_mass(self, volume, density):
        return volume * density


# Создание главного окна
window = tk.Tk()
window.title("Геометрические тела")
window.geometry("400x300")

# Обработчик события для кнопки "Рассчитать"
def calculate():
    shape = shape_var.get()
    length = float(length_entry.get())
    width = float(width_entry.get())
    height = float(height_entry.get())
    density = float(density_entry.get())

    if shape == "Параллелепипед":
        shape_obj = Parallelepiped(length, width, height)
    elif shape == "Тетраэдр":
        shape_obj = Tetrahedron(length, height)
    elif shape == "Шар":
        shape_obj = Sphere(length)

    volume = shape_obj.calculate_volume()
    surface_area = shape_obj.calculate_surface_area()
    mass = shape_obj.calculate_mass(volume, density)

    volume_label.config(text="Объем: {:.2f}".format(volume))
    surface_area_label.config(text="Площадь поверхности: {:.2f}".format(surface_area))
    mass_label.config(text="Масса: {:.2f}".format(mass))


# Обработчик события для кнопки "Сохранить отчет"
def save_report():
    shape = shape_var.get()
    volume = volume_label.cget("text").split(":")[1].strip()
    surface_area = surface_area_label.cget("text").split(":")[1].strip()
    mass = mass_label.cget("text").split(":")[1].strip()

    filetypes = [("Word документ", "*.docx"), ("Excel документ", "*.xlsx")]
    filepath = asksaveasfilename(filetypes=filetypes, defaultextension=filetypes[0][1])

    if shape == "Параллелепипед":
        report = Document()
        report.add_paragraph("Параметры параллелепипеда:")
        report.add_paragraph("Длина: {:.2f}".format(float(length_entry.get())))
        report.add_paragraph("Ширина: {:.2f}".format(float(width_entry.get())))
        report.add_paragraph("Высота: {:.2f}".format(float(height_entry.get())))
        report.add_paragraph("Объем: " + volume)
        report.add_paragraph("Площадь поверхности: " + surface_area)
        report.add_paragraph("Масса: " + mass)
        report.save(filepath)
    elif shape == "Тетраэдр" or shape == "Шар":
        wb = Workbook()
        ws = wb.active
        ws.title = "Результаты расчета"

        ws["A1"] = "Параметры"
        ws["B1"] = "Значения"

        ws["A2"] = "Форма"
        ws["B2"] = shape

        ws["A3"] = "Объем"
        ws["B3"] = volume

        ws["A4"] = "Площадь поверхности"
        ws["B4"] = surface_area

        ws["A5"] = "Масса"
        ws["B5"] = mass

        wb.save(filepath)


# Создание элементов интерфейса
shape_var = tk.StringVar()
shape_var.set("Параллелепипед")
shape_label = tk.Label(window, text="Форма:")
shape_label.pack()
shape_optionmenu = tk.OptionMenu(window, shape_var, "Параллелепипед", "Тетраэдр", "Шар")
shape_optionmenu.pack()

length_label = tk.Label(window, text="Длина:")
length_label.pack()
length_entry = tk.Entry(window)
length_entry.pack()

width_label = tk.Label(window, text="Ширина:")
width_label.pack()
width_entry = tk.Entry(window)
width_entry.pack()

height_label = tk.Label(window, text="Высота:")
height_label.pack()
height_entry = tk.Entry(window)
height_entry.pack()

density_label = tk.Label(window, text="Плотность:")
density_label.pack()
density_entry = tk.Entry(window)
density_entry.pack()

calculate_button = tk.Button(window, text="Рассчитать", command=calculate)
calculate_button.pack()

volume_label = tk.Label(window, text="Объем: ")
volume_label.pack()

surface_area_label = tk.Label(window, text="Площадь поверхности: ")
surface_area_label.pack()

mass_label = tk.Label(window, text="Масса: ")
mass_label.pack()

save_button = tk.Button(window, text="Сохранить отчет", command=save_report)
save_button.pack()

# Запуск главного цикла обработки событий
window.mainloop()
